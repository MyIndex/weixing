#-*- coding:utf-8 -*-
'''
inputDialog
'''
__author__ = 'MoneyIndex'

from PyQt5.QtWidgets import QApplication, QWidget, QLineEdit, QInputDialog, QGridLayout, QLabel, QPushButton, QFrame
import sys
import json
import requests
import itchat
from itchat.content import *
import re
import time
import threading
from docx import Document


# class MyThread(threading.Thread):
#     def __init__(self,fun):
#         super().__init__()
#         self.fun = fun
#         self.singal = threading.Event()
#         self.singal.set()
#
#     def run(self):
#         self.singal.wait()
#         self.fun()
#
#     def pause(self):
#         if self.singal.isSet():
#             self.singal.clear()
#     def restar(self):
#         self.singal.set()


class InputDialog(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()
    def initUI(self):
        self.setWindowTitle('微信服务')
        self.setGeometry(400,400,300,260)

        self.style_one_button = QPushButton('定时发送',self)
        self.style_two_button = QPushButton('自动回复',self)
        self.style_one_button.move(100,50)
        self.style_two_button.move(100,150)

        self.show()


class DSend(QWidget):
    def __init__(self):       
        super().__init__()
        self.initUi()


    @staticmethod
    def is_pict(p, document, username):
        root = str(p)
        pictr_str = '<w:drawing>'
        pictr = re.findall(pictr_str, root)
        if len(pictr) > 0:
            math = re.search('r:embed="(.*)"', root)
            rIds = math.group(1)
            img = document.part.related_parts[rIds]
            with open(r'/Users/qianqibei/Desktop/tmp.jpg', 'wb') as fp:
                fp.write(img.blob)
            pa = r'/Users/qianqibei/Desktop/tmp.jpg'
            itchat.send_image(pa, toUserName=username)
            time.sleep(10)
            return True
        else:
            return False

    @classmethod
    def send_onegroup(cls,name,path):
        itchat.dump_login_status()
        rooms = itchat.get_chatrooms(update=True)
        rooms = itchat.search_chatrooms(name)
        print(rooms)
        print(name)
        if rooms is not None:
            username = rooms[0]['UserName']
            document = Document(path)
            print(document)
            all_num = len(document.paragraphs)
            print(all_num)
            for paragraph in document.paragraphs:
                p = paragraph._element.xml
                if DSend.is_pict(p, document, username):
                    continue
                contant = paragraph.text
                t = len(contant) * 0.3
                itchat.send(contant, toUserName=username)
                time.sleep(t)

        else:
            print('None found')

    def initUi(self):
        self.setWindowTitle("定时发送")
        self.setGeometry(400,400,300,260)

        label1=QLabel("群名:")
        label2=QLabel("发送时间:")
        label3=QLabel("测试时间:")
        label4=QLabel("文件地址:")

        self.nameLable = QLabel("测试")
        self.nameLable.setFrameStyle(QFrame.Panel|QFrame.Sunken)
        self.sendLable = QLabel("07:30")
        self.sendLable.setFrameStyle(QFrame.Panel|QFrame.Sunken)
        self.testLable = QLabel("20:00")
        self.testLable.setFrameStyle(QFrame.Panel|QFrame.Sunken)
        self.addLable = QLabel(r'/Users/qianqibei/Desktop/190416.docx')
        self.addLable.setFrameStyle(QFrame.Panel|QFrame.Sunken)

        self.showLable = QLabel('待运行。。。')
        self.showLable.setFrameStyle(QFrame.Panel|QFrame.Sunken)

        nameButton=QPushButton("...")
        nameButton.clicked.connect(self.selectName)
        sendButton=QPushButton("...")
        sendButton.clicked.connect(self.selectStyle)
        testButton=QPushButton("...")
        testButton.clicked.connect(self.selectNumber)
        addButton=QPushButton("...")
        addButton.clicked.connect(self.selectCost)
        
        TButton = QPushButton('确定')
        PButton = QPushButton('关闭')
        # RButton = QPushButton('继续')
        TButton.clicked.connect(self.Tbutton)
        PButton.clicked.connect(self.Ebutton)
        # RButton.clicked.connect(self.Rbutton)

        mainLayout=QGridLayout()
        mainLayout.addWidget(label1,0,0)
        mainLayout.addWidget(self.nameLable,0,1)
        mainLayout.addWidget(nameButton,0,2)
        mainLayout.addWidget(label2,1,0)
        mainLayout.addWidget(self.sendLable,1,1)
        mainLayout.addWidget(sendButton,1,2)
        mainLayout.addWidget(label3,2,0)
        mainLayout.addWidget(self.testLable,2,1)
        mainLayout.addWidget(testButton,2,2)
        mainLayout.addWidget(label4,3,0)
        mainLayout.addWidget(self.addLable,3,1)
        mainLayout.addWidget(addButton,3,2)
        mainLayout.addWidget(TButton,4,0)
        mainLayout.addWidget(PButton,4,2)
        # mainLayout.addWidget(RButton,4,2)
        mainLayout.addWidget(self.showLable,5,1)

        self.setLayout(mainLayout)



    def selectName(self):
        name,ok = QInputDialog.getText(self,"群名","输入群名:",
                                       QLineEdit.Normal,self.nameLable.text())
        if ok and (len(name)!=0):
            self.nameLable.setText(name)
    def selectStyle(self):
        style,ok = QInputDialog.getText(self,"发送时间","请选择发送时间：",QLineEdit.Normal,self.sendLable.text())
        if ok and (len(style) != 0):
            self.sendLable.setText(style)

    def selectNumber(self):
        number,ok = QInputDialog.getText(self,"测试时间","请输入测试时间：",QLineEdit.Normal,self.testLable.text())
        if ok and (len(number) != 0):
            self.testLable.setText(str(number))

    def selectCost(self):
        cost,ok = QInputDialog.getText(self,"文件地址","请输入文件地址：",QLineEdit.Normal,self.addLable.text())
        if ok and (len(cost) != 0):
            self.addLable.setText(str(cost))

    def send(self):
        itchat.auto_login()
        while (1):
            date = time.strftime('%y%m%d %H:%M:%S')
            if date[7:12] == self.sendLable.text() or date[7:12] == self.testLable.text():
                DSend.send_onegroup(self.nameLable.text(),self.addLable.text())
            time.sleep(60)
    def Tbutton(self):
        self.thread = threading.Thread(target=self.send)
        self.showLable.setText('开始')
        self.thread.start()

    def Ebutton(self):
        raise ValueError
        #     self.showLable.setText('暂停')


    # def Rbutton(self):
    #     self.showLable.setText('继续')
    #     self.thread.restar()

class Recall(QWidget):
    def __init__(self):
        super().__init__()
        self.initUi()


    @staticmethod
    def tuling(text):
        try:
            url = 'http://openapi.tuling123.com/openapi/api/v2'
            postdata = {"perception": {
                "inputText": {
                    "text": text
                },
            },
                "userInfo": {
                    "apiKey": "06decbd2f80348dc9d6a381d38784938",
                    "userId": "429075"
                }
            }
            postdata = json.dumps(postdata)
            req = requests.post(url, data=postdata)
            data = req.json()
            contant = data['results'][0]['values']['text']
            return contant
        except Exception as e:
            print(e)
            return 'hello'

    @itchat.msg_register(TEXT, isGroupChat=True)
    def recall(msg):
        itchat.dump_login_status()
        myname = itchat.get_friends(update=True)[0]['UserName']
        rooms = itchat.search_chatrooms(name=room_name)[0]['UserName']
        username = msg['FromUserName']
        if username == rooms:
            if msg['isAt']:
                text = msg['Text']
                name = msg['ActualNickName']
                contant = Recall.tuling(text)
                contant += '@%s ' % name
                itchat.send(contant, username)

    def initUi(self):
        self.setWindowTitle('自动回复')
        self.setGeometry(400, 400, 300, 260)

        label1 = QLabel('群名:')
        self.nameLable = QLabel('测试：')
        self.nameLable.setFrameStyle(QFrame.Panel | QFrame.Sunken)
        self.showLable = QLabel('待运行。。。')
        self.showLable.setFrameStyle(QFrame.Panel | QFrame.Sunken)

        nameButton = QPushButton('....')
        nameButton.clicked.connect(self.selectName)

        TButton = QPushButton('确定')
        PButton = QPushButton('关闭')
        # RButton = QPushButton('继续')
        TButton.clicked.connect(self.Tbutton)
        PButton.clicked.connect(self.Ebutton)
        # RButton.clicked.connect(self.Rbutton)

        mainLayout = QGridLayout()
        mainLayout.addWidget(label1,0,0)
        mainLayout.addWidget(self.nameLable,0,1)
        mainLayout.addWidget(nameButton,0,2)


        mainLayout.addWidget(TButton, 1, 0)
        mainLayout.addWidget(PButton, 1, 2)
        # mainLayout.addWidget(RButton, 1, 2)
        mainLayout.addWidget(self.showLable, 2, 1)

        self.setLayout(mainLayout)

    def selectName(self):
        name,ok = QInputDialog.getText(self,'群名','输入群名',QLineEdit.Normal,self.nameLable.text())
        if ok and (len(name)!=0):
            self.nameLable.setText(name)
    def call(self):
        itchat.auto_login()
        global room_name
        room_name = self.nameLable.text()
        itchat.run()

    def Tbutton(self):
        self.thread = threading.Thread(target=self.call())
        self.showLable.setText('开始')
        self.thread.start()

    def Ebutton(self):

        raise None

    # def Rbutton(self):
    #     self.showLable.setText('继续')
    #     self.thread.restar()


if __name__=="__main__":
    # t = threading.Thread(target=itchat.auto_login)
    # t.start()###登录，扫码，相当于登录微信网页版
    app=QApplication(sys.argv)
    myshow=InputDialog()
    dsend = DSend()
    recall = Recall()
    myshow.show()
    myshow.style_one_button.clicked.connect(dsend.show)
    myshow.style_two_button.clicked.connect(recall.show)
    sys.exit(app.exec_())
